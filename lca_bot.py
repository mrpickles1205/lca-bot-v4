
import streamlit as st
import pandas as pd
import random
from docx import Document
from docx.shared import Inches
import datetime

def simulate_web_summary(product):
    return {
        "materials": f"The {product} typically includes plastic casing (polypropylene), metal parts (stainless steel, copper), lithium-ion battery, and small electronics.",
        "manufacturing": f"Manufacturing involves injection molding for plastic components, PCB assembly, battery installation, and final casing assembly in factories primarily located in Asia.",
        "use_phase": f"During use, the {product} consumes around 3 watts per charge, typically charged once per week, with an average lifetime of 3 years.",
        "end_of_life": f"The {product} is partially recyclable, but batteries and electronics often end up in landfill or incineration due to complex disassembly requirements."
    }

def create_docx_report(product, summary):
    doc = Document()
    doc.add_heading(f'LCA Report: {product}', 0)
    doc.add_paragraph(f"Date: {datetime.date.today()}")

    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"This LCA report presents a cradle-to-grave assessment of the {product}. The analysis includes material sourcing, manufacturing, use phase, and end-of-life considerations based on publicly available sources.")

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(f"The goal of this study is to evaluate the environmental impact of one {product} across its life cycle, aligned with ISO 14040 and 14044.")

    doc.add_heading("2. Materials & Composition", level=1)
    doc.add_paragraph(summary["materials"])

    doc.add_heading("3. Manufacturing Process", level=1)
    doc.add_paragraph(summary["manufacturing"])

    doc.add_heading("4. Use Phase", level=1)
    doc.add_paragraph(summary["use_phase"])

    doc.add_heading("5. End-of-Life Management", level=1)
    doc.add_paragraph(summary["end_of_life"])

    doc.add_heading("6. Life Cycle Inventory (LCI)", level=1)
    df = pd.DataFrame({
        'Phase': ['Materials', 'Manufacturing', 'Use Phase', 'End-of-Life'],
        'Energy (MJ)': [random.uniform(20, 100) for _ in range(4)],
        'GHG Emissions (kg CO2-eq)': [random.uniform(2, 20) for _ in range(4)],
        'Water Use (L)': [random.uniform(10, 50) for _ in range(4)]
    })
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(round(val, 2)) if isinstance(val, (int, float)) else str(val)

    doc.add_heading("7. Life Cycle Impact Assessment (LCIA)", level=1)
    doc.add_paragraph("Impact categories include GWP (Global Warming Potential), cumulative energy demand, and water footprint.")

    doc.add_heading("8. Interpretation", level=1)
    doc.add_paragraph(f"The largest contributor to GHG emissions is likely the manufacturing or use phase depending on the device's energy consumption and frequency of use.")

    doc.add_heading("Appendix A: Definitions", level=2)
    doc.add_paragraph("GHG: Greenhouse Gases. CO2-eq: Carbon dioxide equivalent. LCA: Life Cycle Assessment.")

    file_path = f"ISO_LCA_{product.replace(' ', '_')}.docx"
    doc.save(file_path)
    return file_path

st.title("🌍 AI-Powered ISO LCA Bot")
product_name = st.text_input("Enter a product to assess:", "Electric Toothbrush")

if st.button("Generate Full LCA Report"):
    with st.spinner("Gathering data and writing report..."):
        summary = simulate_web_summary(product_name)
        file_path = create_docx_report(product_name, summary)

    with open(file_path, "rb") as f:
        st.download_button("📄 Download Detailed DOCX Report", f, file_name=file_path)
